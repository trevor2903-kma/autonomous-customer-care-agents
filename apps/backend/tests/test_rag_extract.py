"""Phase 2 — extract_text (đa định dạng) + chunk_text (chunking tổng quát). Offline, KHÔNG network."""

from __future__ import annotations

from io import BytesIO

import pytest

from app.services.extract import extract_text
from app.services.rag_service import chunk_text


def test_extract_txt_and_md() -> None:
    assert extract_text("a.txt", "xin chào shop".encode("utf-8")) == "xin chào shop"
    assert extract_text("b.md", "# tiêu đề\nnội dung".encode("utf-8")) == "# tiêu đề\nnội dung"


def test_extract_docx_reads_paragraphs() -> None:
    from docx import Document

    doc = Document()
    doc.add_paragraph("Chính sách đổi trả trong 7 ngày.")
    doc.add_paragraph("Phí ship nội thành 20k.")
    buf = BytesIO()
    doc.save(buf)

    text = extract_text("policy.docx", buf.getvalue())
    assert "đổi trả" in text
    assert "Phí ship" in text


def test_extract_unsupported_suffix_raises() -> None:
    with pytest.raises(ValueError):
        extract_text("image.png", b"\x89PNG...")


def test_chunk_text_generic_multiple_chunks() -> None:
    text = " ".join(
        f"Câu số {i} nói về chính sách của shop quần áo, viết dài để đủ ký tự." for i in range(60)
    )
    chunks = chunk_text(text, size=300, overlap=60)
    assert len(chunks) > 1
    assert all(isinstance(c, str) and c.strip() for c in chunks)


def test_chunk_text_empty_returns_empty() -> None:
    assert chunk_text("   \n\n  ") == []
